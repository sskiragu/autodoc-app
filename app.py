# app.py
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
import io

def extract_text_from_docx(file):
    doc = Document(file)
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_from_pdf(file):
    text = ""
    reader = PyPDF2.PdfReader(file)
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def create_modern_cv(text):
    buffer = io.BytesIO()
    doc = Document()

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if not lines:
        return None

    # ======= Name Section =======
    name = lines[0]
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(22)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ======= Contact Info =======
    for i in range(1, 3):  # next 2 lines
        p = doc.add_paragraph(lines[i])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # empty line

    # ======= Body Sections =======
    #==============================#
    sections = {
        "summary": [],
        "experience": [],
        "education": [],
        "skills": []
    }

    current = None
    for line in lines[3:]:
        lower = line.lower()
        if "experience" in lower:
            current = "experience"
            continue
        elif "education" in lower:
            current = "education"
            continue
        elif "skills" in lower:
            current = "skills"
            continue
        elif current is None:
            current = "summary"

        if current:
            sections[current].append(line)

    def add_section(title, content, bullet=False):
        if content:
            doc.add_heading(title, level=2)
            for line in content:
                if bullet:
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    doc.add_paragraph(line)

    add_section("Summary", sections["summary"])
    add_section("Experience", sections["experience"])
    add_section("Education", sections["education"])
    add_section("Skills", sections["skills"], bullet=True)

    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.title("🧠 AutoDoc CV Generator")
uploaded_file = st.file_uploader("Upload your CV (DOCX or PDF)", type=["docx", "pdf"])

if uploaded_file:
    file_type = uploaded_file.name.split('.')[-1].lower()

    if file_type == "docx":
        text = extract_text_from_docx(uploaded_file)
    elif file_type == "pdf":
        text = extract_text_from_pdf(uploaded_file)
    else:
        st.error("Unsupported file type.")
        text = None

    if text:
        st.subheader("Extracted Text")
        st.text_area("Preview", text, height=250)

        buffer = create_modern_cv(text)
        if buffer:
            st.download_button(
                "⬇️ Download Modern CV",
                data=buffer,
                file_name="modern_cv.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Could not generate CV.")
